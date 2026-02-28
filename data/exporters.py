"""Export literature review results to BibTeX, RIS, CSV, DOCX, and JSON."""
from __future__ import annotations

import io
import json
import re
from datetime import datetime

import pandas as pd


def _clean(text: str | None) -> str:
    if not text:
        return ""
    return re.sub(r"\s+", " ", str(text)).strip()


def _author_list(authors_json) -> list[str]:
    if not authors_json:
        return []
    try:
        authors = json.loads(authors_json) if isinstance(authors_json, str) else authors_json
        return [a.get("name", "") for a in authors if a.get("name")]
    except Exception:
        return []


def papers_to_bibtex(papers: list[dict]) -> str:
    """Return a BibTeX string for the given papers."""
    lines: list[str] = []
    for p in papers:
        authors = _author_list(p.get("authors"))
        key_parts = []
        if authors:
            last = authors[0].split()[-1] if authors[0].split() else "Unknown"
            key_parts.append(last)
        if p.get("year"):
            key_parts.append(str(p["year"]))
        title_words = (_clean(p.get("title")) or "").split()[:2]
        key_parts.extend(title_words)
        key = re.sub(r"[^a-zA-Z0-9]", "", "".join(key_parts))[:30] or p["id"][:12]

        doc_type = p.get("document_type", "article")
        bib_type = "article" if doc_type in ("article", "review") else "inproceedings" if "conference" in doc_type else "misc"

        lines.append(f"@{bib_type}{{{key},")
        lines.append(f'  title = {{{_clean(p.get("title"))}}},')
        if authors:
            lines.append(f'  author = {{{" and ".join(authors)}}},')
        if p.get("year"):
            lines.append(f'  year = {{{p["year"]}}},')
        if p.get("journal"):
            lines.append(f'  journal = {{{_clean(p.get("journal"))}}},')
        if p.get("doi"):
            lines.append(f'  doi = {{{p["doi"]}}},')
        if p.get("abstract"):
            lines.append(f'  abstract = {{{_clean(p.get("abstract"))[:400]}}},')
        lines.append("}")
        lines.append("")
    return "\n".join(lines)


def papers_to_ris(papers: list[dict]) -> str:
    """Return an RIS-format string."""
    lines: list[str] = []
    doc_type_map = {
        "article": "JOUR",
        "review": "JOUR",
        "conference": "CONF",
        "book-chapter": "CHAP",
        "preprint": "JOUR",
    }
    for p in papers:
        ty = doc_type_map.get(p.get("document_type", ""), "JOUR")
        lines.append(f"TY  - {ty}")
        if p.get("title"):
            lines.append(f"TI  - {_clean(p['title'])}")
        for a in _author_list(p.get("authors")):
            lines.append(f"AU  - {a}")
        if p.get("year"):
            lines.append(f"PY  - {p['year']}")
        if p.get("journal"):
            lines.append(f"JO  - {_clean(p['journal'])}")
        if p.get("doi"):
            lines.append(f"DO  - {p['doi']}")
        if p.get("abstract"):
            lines.append(f"AB  - {_clean(p.get('abstract'))[:500]}")
        if p.get("open_access_url"):
            lines.append(f"UR  - {p['open_access_url']}")
        lines.append("ER  - ")
        lines.append("")
    return "\n".join(lines)


def papers_to_dataframe(papers: list[dict]) -> pd.DataFrame:
    """Return a clean DataFrame of included papers."""
    rows = []
    for p in papers:
        authors = _author_list(p.get("authors"))
        rows.append({
            "Title": _clean(p.get("title")),
            "Authors": "; ".join(authors[:5]),
            "Year": p.get("year"),
            "Journal": _clean(p.get("journal")),
            "DOI": p.get("doi"),
            "Citation Count": p.get("citation_count"),
            "Quality Score": p.get("quality_score"),
            "Cluster": p.get("cluster_label"),
            "Inclusion Decision": p.get("final_status"),
            "Screening Reason": p.get("screening_pass1_reason"),
            "Quality Notes": p.get("quality_notes"),
            "Open Access URL": p.get("open_access_url"),
            "Abstract": _clean(p.get("abstract")),
        })
    return pd.DataFrame(rows)


def papers_to_docx_bytes(papers: list[dict], synthesis: dict | None) -> bytes:
    """Return a DOCX report as bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading("Literature Review Report", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Papers included: {len(papers)}")

    if synthesis:
        doc.add_heading("Literature Synthesis", 1)

        if synthesis.get("narrative_overview"):
            doc.add_heading("Overview", 2)
            doc.add_paragraph(synthesis["narrative_overview"])

        if synthesis.get("key_themes"):
            doc.add_heading("Key Themes", 2)
            for theme in synthesis["key_themes"]:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(theme)

        if synthesis.get("research_gaps"):
            doc.add_heading("Research Gaps", 2)
            for gap in synthesis["research_gaps"]:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(gap)

        if synthesis.get("key_debates"):
            doc.add_heading("Key Debates", 2)
            for debate in synthesis["key_debates"]:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(debate)

    doc.add_heading("Included Papers", 1)
    for i, p in enumerate(papers, 1):
        authors = _author_list(p.get("authors"))
        heading = f"{i}. {_clean(p.get('title'))}"
        doc.add_heading(heading, 3)
        meta = []
        if authors:
            meta.append("; ".join(authors[:3]))
        if p.get("year"):
            meta.append(str(p["year"]))
        if p.get("journal"):
            meta.append(_clean(p["journal"]))
        if meta:
            doc.add_paragraph(" | ".join(meta))
        if p.get("abstract"):
            doc.add_paragraph(_clean(p["abstract"])[:600] + "…")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_audit_trail(log_entries: list[dict], papers: list[dict]) -> str:
    """Return a JSON audit trail."""
    trail = {
        "generated_at": datetime.now().isoformat(),
        "pipeline_log": log_entries,
        "paper_decisions": [
            {
                "id": p["id"],
                "title": p.get("title"),
                "screening_pass1": p.get("screening_pass1"),
                "screening_pass1_reason": p.get("screening_pass1_reason"),
                "human_decision": p.get("human_decision"),
                "quality_score": p.get("quality_score"),
                "final_status": p.get("final_status"),
            }
            for p in papers
        ],
    }
    return json.dumps(trail, indent=2, default=str)
